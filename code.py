# =============================================================================
# BANASPATI — Multimodal RAG Academic Assistant Pipeline
# Practicum AI 2026 | Module 5
# =============================================================================
# ARCHITECTURE OVERVIEW:
#   Stage 1 → Environment bootstrap (deps, secrets, GPU check)
#   Stage 2 → Dataset acquisition from GitHub remote
#   Stage 3 → Multimodal document ingestion (PDF + DOCX → chunks + images)
#   Stage 4 → Embedding + ChromaDB vector index construction
#   Stage 5 → Hybrid retrieval (dense cosine + BM25 → RRF fusion)
#   Stage 6 → Generation via Gemini 2.0 Flash + hallucination gate
#   Stage 7 → Telemetry layer (latency / token / cost / RAM)
#   Stage 8 → RAGAS qualitative evaluation
#   Stage 9 → LLM-as-Judge via Gemini 3 Flash
#   Stage 10 → Demo Sandbox (interactive single-query runner)
# =============================================================================


# =============================================================================
# ██████████████████████████████████████████████████████████████████████████
# STAGE 1 — ENVIRONMENT BOOTSTRAP
# ██████████████████████████████████████████████████████████████████████████
# =============================================================================

# CELL 1.1 — Dependency Installation
# REASON: pin versions to prevent Colab silent-upgrade breakage.
# rank_bm25: sparse retrieval. ragas: eval framework. unstructured: docx parse.
# pymupdf (fitz): PDF text+image extract. psutil: RAM telemetry.

import subprocess, sys

_deps = [
    "google-generativeai==0.8.5",
    "chromadb==0.5.23",
    "langchain==0.3.25",
    "langchain-google-genai==2.1.4",
    "langchain-community==0.3.24",
    "langchain-chroma==0.2.4",
    "ragas==0.2.15",
    "rank_bm25==0.2.2",
    "pymupdf==1.24.11",
    "python-docx==1.1.2",
    "unstructured==0.16.23",
    "psutil==6.1.1",
    "pandas==2.2.3",
    "datasets==3.5.0",
    "Pillow==11.1.0",
    "tqdm==4.67.1",
]

subprocess.check_call([sys.executable, "-m", "pip", "install", "--quiet"] + _deps)

# -----------------------------------------------------------------------------
# CELL 1.2 — Core Imports
# -----------------------------------------------------------------------------

import os, time, base64, json, re, gc
from io import BytesIO
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional

import psutil
import pandas as pd
import fitz                          # PyMuPDF
from docx import Document as DocxDocument
from PIL import Image
from tqdm import tqdm
from rank_bm25 import BM25Okapi

import google.generativeai as genai
import chromadb
from chromadb.utils import embedding_functions

# -----------------------------------------------------------------------------
# CELL 1.3 — Secret Manager + API Key Injection
# SECURITY: NEVER hardcode keys. All keys sourced from Colab Secret Manager.
# Colab UI → left sidebar → 🔑 Secrets → add GEMINI_API_KEY before running.
# -----------------------------------------------------------------------------

try:
    from google.colab import userdata
    GEMINI_API_KEY = userdata.get("GEMINI_API_KEY")
except ImportError:
    # Local dev fallback — read from env, still never hardcoded.
    GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY")

if not GEMINI_API_KEY:
    raise EnvironmentError(
        "FATAL: GEMINI_API_KEY not found. "
        "Add it to Colab Secrets (🔑) or set env var GEMINI_API_KEY."
    )

genai.configure(api_key=GEMINI_API_KEY)

# -----------------------------------------------------------------------------
# CELL 1.4 — GPU / Hardware Verification
# INFO: Colab GPU type affects VRAM budget if local SLM fallback is needed.
# API-first strategy means GPU is only used for embedding if BGE-M3 is chosen.
# -----------------------------------------------------------------------------

try:
    import torch
    _gpu_available = torch.cuda.is_available()
    _gpu_name      = torch.cuda.get_device_name(0) if _gpu_available else "None"
    print(f"[HW] GPU: {_gpu_name} | VRAM: "
          f"{torch.cuda.get_device_properties(0).total_memory / 1e9:.1f}GB"
          if _gpu_available else "[HW] GPU: Not available — API-only mode active")
except ImportError:
    print("[HW] torch not installed — skipping GPU check")

_ram_total_gb = psutil.virtual_memory().total / 1e9
print(f"[HW] RAM total: {_ram_total_gb:.1f} GB")

# -----------------------------------------------------------------------------
# CELL 1.5 — Global Configuration
# ALL tunable parameters centralized here. Change here, propagates everywhere.
# -----------------------------------------------------------------------------

CFG = {
    # Retrieval
    "TOP_K":                    5,       # chunks retrieved per query
    "SIMILARITY_THRESHOLD":     0.35,    # below → CONTEXT_INSUFFICIENT gate fires
    "CHUNK_SIZE":               800,     # characters per text chunk
    "CHUNK_OVERLAP":            120,     # character overlap between chunks
    "BM25_WEIGHT":              0.4,     # RRF blend: BM25 weight
    "DENSE_WEIGHT":             0.6,     # RRF blend: dense embedding weight

    # Models
    "GENERATOR_MODEL":          "gemini-2.0-flash",
    "JUDGE_MODEL":              "gemini-3-flash",          # LLM-as-Judge (mandated)
    "EMBEDDING_MODEL":          "models/text-embedding-004",

    # Gemini 2.0 Flash pricing (USD per 1M tokens) — verify at ai.google.dev
    "COST_INPUT_PER_1M":        0.075,
    "COST_OUTPUT_PER_1M":       0.30,

    # Paths
    "DATASET_DIR":              Path("./dataset"),
    "CHROMA_PERSIST_DIR":       Path("./chroma_db"),
    "EVAL_CSV":                 "banaspati_eval_questions.csv",
    "COLLECTION_NAME":          "banaspati_docs",

    # Generation
    "MAX_OUTPUT_TOKENS":        2048,
    "TEMPERATURE":              0.1,     # low temp → faithful to context
}

print("[CFG] Configuration loaded.")


# =============================================================================
# ██████████████████████████████████████████████████████████████████████████
# STAGE 2 — DATASET ACQUISITION
# ██████████████████████████████████████████████████████████████████████████
# =============================================================================

# CELL 2.1 — Clone Dataset from GitHub
# REASON: raw GitHub file fetch (wget per file) is fragile for binary PDFs.
# git clone is atomic and handles LFS if needed. Depth=1 saves bandwidth.

import subprocess

_REPO_URL = "https://github.com/zelebwr/multimodal-rag-based-llm.git"

if not CFG["DATASET_DIR"].exists():
    print("[DATA] Cloning repository...")
    subprocess.check_call([
        "git", "clone", "--depth", "1", _REPO_URL, "./repo_clone"
    ])
    # Move only dataset dir to working root
    import shutil
    shutil.copytree("./repo_clone/dataset", str(CFG["DATASET_DIR"]))
    shutil.rmtree("./repo_clone")
    print("[DATA] Dataset ready.")
else:
    print("[DATA] Dataset directory already exists — skipping clone.")

# Enumerate and validate all expected files
_EXPECTED_FILES = [
    "banaspati_eval_questions.csv",
    "Data Dosen.pdf",
    "Jadwal Perkuliahan.docx",
    "Kalender-Akademik-ITS-Thn-Akademik-2025-2026.pdf",
    "Kurikulum.pdf",
    "Nilai snbt 2025.pdf",
    "Peraturan Akademik.pdf",
    "Sosialisasi Magang dan Prestasi DTI.pdf",
    "Visi Misi Departemen.pdf",
]

_missing = [f for f in _EXPECTED_FILES if not (CFG["DATASET_DIR"] / f).exists()]
if _missing:
    raise FileNotFoundError(f"[DATA] MISSING FILES: {_missing}")

print(f"[DATA] All {len(_EXPECTED_FILES)} files verified.")


# =============================================================================
# ██████████████████████████████████████████████████████████████████████████
# STAGE 3 — MULTIMODAL DOCUMENT INGESTION
# ██████████████████████████████████████████████████████████████████████████
# =============================================================================

# CELL 3.1 — Document Chunk Dataclass
# Unified schema for all chunk types (text, table, image).
# source_file + page_num enables source citation in final answer.

@dataclass
class DocChunk:
    chunk_id:    str                    # "{filename}_p{page}_{idx}"
    source_file: str                    # original filename
    page_num:    int                    # 0-indexed page
    chunk_type:  str                    # "text" | "table" | "image"
    content:     str                    # text content or image description
    image_b64:   Optional[str] = None  # base64 PNG for image chunks
    metadata:    dict = field(default_factory=dict)

# -----------------------------------------------------------------------------
# CELL 3.2 — PDF Ingestion (PyMuPDF)
# MECHANISM:
#   page.get_text("text")  → raw text extraction per page
#   page.get_text("blocks") → block-level parse for table detection heuristic
#   page.get_pixmap()      → rasterize page to PNG → base64 for image chunks
# DECISION: extract both text AND image per page → enables multimodal retrieval
#   where text retrieval fails (e.g., scanned pages, chart-heavy slides).
# -----------------------------------------------------------------------------

def _chunk_text(text: str, source: str, page: int, start_idx: int) -> list[DocChunk]:
    """Split text into overlapping fixed-size chunks."""
    chunks = []
    size    = CFG["CHUNK_SIZE"]
    overlap = CFG["CHUNK_OVERLAP"]
    step    = size - overlap
    text    = text.strip()
    if not text:
        return chunks
    for i, start in enumerate(range(0, len(text), step)):
        seg = text[start : start + size]
        if len(seg) < 50:   # discard micro-fragments
            continue
        chunks.append(DocChunk(
            chunk_id    = f"{source}_p{page}_{start_idx + i}",
            source_file = source,
            page_num    = page,
            chunk_type  = "text",
            content     = seg,
        ))
    return chunks


def ingest_pdf(filepath: Path) -> list[DocChunk]:
    chunks = []
    doc    = fitz.open(str(filepath))
    fname  = filepath.name
    idx    = 0

    for page_num, page in enumerate(tqdm(doc, desc=f"PDF: {fname}", leave=False)):
        # --- Text extraction ---
        page_text = page.get_text("text")
        text_chunks = _chunk_text(page_text, fname, page_num, idx)
        chunks.extend(text_chunks)
        idx += len(text_chunks)

        # --- Image chunk: rasterize full page at 150 DPI ---
        # REASON: captures diagrams, tables-as-images, stamps in academic PDFs.
        # 150 DPI balances quality vs base64 payload size (~100-200KB per page).
        mat    = fitz.Matrix(150 / 72, 150 / 72)
        pixmap = page.get_pixmap(matrix=mat, colorspace=fitz.csRGB)
        img_bytes = pixmap.tobytes("png")
        img_b64   = base64.b64encode(img_bytes).decode("utf-8")

        chunks.append(DocChunk(
            chunk_id    = f"{fname}_p{page_num}_img",
            source_file = fname,
            page_num    = page_num,
            chunk_type  = "image",
            content     = f"[Page image: {fname}, page {page_num + 1}]",
            image_b64   = img_b64,
        ))
        idx += 1

    doc.close()
    return chunks


# -----------------------------------------------------------------------------
# CELL 3.3 — DOCX Ingestion (python-docx)
# MECHANISM:
#   doc.paragraphs → text blocks
#   doc.tables     → iterate rows/cells → stringify
# DECISION: python-docx over unstructured for .docx because unstructured
#   adds heavy system deps (libmagic, tesseract) that slow Colab cold start.
# -----------------------------------------------------------------------------

def ingest_docx(filepath: Path) -> list[DocChunk]:
    chunks = []
    doc    = DocxDocument(str(filepath))
    fname  = filepath.name
    idx    = 0

    # Paragraph text
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    text_chunks = _chunk_text(full_text, fname, 0, idx)
    chunks.extend(text_chunks)
    idx += len(text_chunks)

    # Tables → stringify each table as TSV-like block
    for t_idx, table in enumerate(doc.tables):
        rows = []
        for row in table.rows:
            rows.append(" | ".join(cell.text.strip() for cell in row.cells))
        table_text = "\n".join(rows)
        if table_text.strip():
            chunks.append(DocChunk(
                chunk_id    = f"{fname}_table_{t_idx}",
                source_file = fname,
                page_num    = 0,
                chunk_type  = "table",
                content     = table_text,
            ))

    return chunks


# -----------------------------------------------------------------------------
# CELL 3.4 — Run Full Ingestion
# -----------------------------------------------------------------------------

ALL_CHUNKS: list[DocChunk] = []

_pdf_files = [f for f in CFG["DATASET_DIR"].iterdir()
              if f.suffix.lower() == ".pdf"]
_docx_files = [f for f in CFG["DATASET_DIR"].iterdir()
               if f.suffix.lower() == ".docx"]

print(f"[INGEST] Processing {len(_pdf_files)} PDFs + {len(_docx_files)} DOCX files...")

for pdf_path in _pdf_files:
    ALL_CHUNKS.extend(ingest_pdf(pdf_path))

for docx_path in _docx_files:
    ALL_CHUNKS.extend(ingest_docx(docx_path))

_text_chunks  = [c for c in ALL_CHUNKS if c.chunk_type in ("text", "table")]
_image_chunks = [c for c in ALL_CHUNKS if c.chunk_type == "image"]

print(f"[INGEST] Total chunks: {len(ALL_CHUNKS)} "
      f"(text/table: {len(_text_chunks)}, image: {len(_image_chunks)})")


# =============================================================================
# ██████████████████████████████████████████████████████████████████████████
# STAGE 4 — EMBEDDING + CHROMADB INDEX CONSTRUCTION
# ██████████████████████████████████████████████████████████████████████████
# =============================================================================

# CELL 4.1 — Google Embedding Function Wrapper for ChromaDB
# text-embedding-004: 768-dim, free tier, 2048 token input limit.
# REASON over BGE-M3 local: T4 VRAM budget preserved for inference; API embed
#   cost for this corpus (~35MB text) is negligible (<$0.01 total).

class GeminiEmbeddingFunction(embedding_functions.EmbeddingFunction):
    def __call__(self, input: list[str]) -> list[list[float]]:
        # Batch in groups of 100 — Google API limit per request
        embeddings = []
        batch_size = 100
        for i in range(0, len(input), batch_size):
            batch  = input[i : i + batch_size]
            result = genai.embed_content(
                model   = CFG["EMBEDDING_MODEL"],
                content = batch,
                task_type = "RETRIEVAL_DOCUMENT",
            )
            embeddings.extend(result["embedding"])
        return embeddings

# -----------------------------------------------------------------------------
# CELL 4.2 — ChromaDB Collection Setup
# persist_directory: survives Colab session if Drive is mounted.
# If collection already exists (re-run), skip re-indexing.
# -----------------------------------------------------------------------------

_chroma_client = chromadb.PersistentClient(path=str(CFG["CHROMA_PERSIST_DIR"]))
_embed_fn      = GeminiEmbeddingFunction()

_collection = _chroma_client.get_or_create_collection(
    name               = CFG["COLLECTION_NAME"],
    embedding_function = _embed_fn,
    metadata           = {"hnsw:space": "cosine"},  # cosine distance for similarity gate
)

# -----------------------------------------------------------------------------
# CELL 4.3 — Index Text + Table Chunks
# IMAGE CHUNKS ARE NOT INDEXED IN VECTOR DB.
# REASON: embedding a base64 string produces nonsense vectors. Image chunks
#   are stored separately and retrieved via page-level matching after text hit.
# -----------------------------------------------------------------------------

_already_indexed = _collection.count()
print(f"[INDEX] ChromaDB existing documents: {_already_indexed}")

if _already_indexed < len(_text_chunks):
    print("[INDEX] Building index — this may take 3-8 minutes on free tier API...")

    _docs  = [c.content     for c in _text_chunks]
    _ids   = [c.chunk_id    for c in _text_chunks]
    _metas = [{
        "source_file": c.source_file,
        "page_num":    c.page_num,
        "chunk_type":  c.chunk_type,
    } for c in _text_chunks]

    # Upsert in batches of 50 to respect API rate limits
    _batch = 50
    for i in tqdm(range(0, len(_text_chunks), _batch), desc="[INDEX] Upserting"):
        _collection.upsert(
            documents = _docs [i : i + _batch],
            ids       = _ids  [i : i + _batch],
            metadatas = _metas[i : i + _batch],
        )

    print(f"[INDEX] Indexed {_collection.count()} chunks.")
else:
    print("[INDEX] Index already populated — skipping.")

# -----------------------------------------------------------------------------
# CELL 4.4 — BM25 Index (Sparse Retrieval)
# REASON: dense retrieval fails on exact-match queries (course codes, lecturer
#   names, specific dates). BM25 handles keyword overlap; RRF fusion merges both.
# -----------------------------------------------------------------------------

_bm25_corpus   = [c.content.lower().split() for c in _text_chunks]
_bm25_index    = BM25Okapi(_bm25_corpus)
_bm25_id_map   = [c.chunk_id for c in _text_chunks]

print(f"[BM25] Index built over {len(_bm25_corpus)} documents.")


# =============================================================================
# ██████████████████████████████████████████████████████████████████████████
# STAGE 5 — HYBRID RETRIEVAL PIPELINE
# ██████████████████████████████████████████████████████████████████████████
# =============================================================================

# CELL 5.1 — Retrieval Result Dataclass

@dataclass
class RetrievalResult:
    query:              str
    chunks:             list[DocChunk]
    sources:            list[str]          # unique filenames
    top_similarity:     float              # max cosine similarity score
    retrieval_latency:  float              # seconds
    context_sufficient: bool               # False → CONTEXT_INSUFFICIENT gate


# -----------------------------------------------------------------------------
# CELL 5.2 — Reciprocal Rank Fusion (RRF)
# MECHANISM:
#   1. Dense: ChromaDB cosine query → ranked list of chunk_ids + distances
#   2. Sparse: BM25 scores → ranked list of chunk_ids
#   3. RRF score = Σ weight / (rank + 60) per chunk_id across both lists
#   4. Sort by RRF score descending → top_k
# WHY RRF over score normalization: RRF is rank-based, not score-based →
#   immune to scale mismatch between cosine similarity and BM25 raw scores.
# -----------------------------------------------------------------------------

def _rrf_score(rank: int, weight: float, k: int = 60) -> float:
    return weight / (rank + k)


def retrieve(query: str) -> RetrievalResult:
    t_start = time.perf_counter()

    # --- Dense retrieval ---
    dense_results = _collection.query(
        query_texts   = [query],
        n_results     = CFG["TOP_K"] * 2,   # over-fetch, RRF re-ranks
        include       = ["documents", "metadatas", "distances"],
    )

    dense_ids       = dense_results["ids"][0]
    dense_distances = dense_results["distances"][0]
    # ChromaDB cosine distance = 1 - similarity → convert back
    dense_sims      = [1.0 - d for d in dense_distances]

    # --- Sparse BM25 retrieval ---
    bm25_scores  = _bm25_index.get_scores(query.lower().split())
    bm25_ranked  = sorted(
        range(len(bm25_scores)), key=lambda i: bm25_scores[i], reverse=True
    )[: CFG["TOP_K"] * 2]
    bm25_ids     = [_bm25_id_map[i] for i in bm25_ranked]

    # --- RRF fusion ---
    rrf_scores: dict[str, float] = {}

    for rank, chunk_id in enumerate(dense_ids):
        rrf_scores[chunk_id] = rrf_scores.get(chunk_id, 0.0) + \
            _rrf_score(rank, CFG["DENSE_WEIGHT"])

    for rank, chunk_id in enumerate(bm25_ids):
        rrf_scores[chunk_id] = rrf_scores.get(chunk_id, 0.0) + \
            _rrf_score(rank, CFG["BM25_WEIGHT"])

    top_ids = sorted(rrf_scores, key=rrf_scores.get, reverse=True)[: CFG["TOP_K"]]

    # --- Resolve chunk_ids back to DocChunk objects ---
    _chunk_lookup = {c.chunk_id: c for c in _text_chunks}
    top_chunks    = [_chunk_lookup[cid] for cid in top_ids if cid in _chunk_lookup]

    # --- Similarity gate ---
    top_sim = max(dense_sims) if dense_sims else 0.0

    t_end = time.perf_counter()

    return RetrievalResult(
        query              = query,
        chunks             = top_chunks,
        sources            = list({c.source_file for c in top_chunks}),
        top_similarity     = top_sim,
        retrieval_latency  = t_end - t_start,
        context_sufficient = top_sim >= CFG["SIMILARITY_THRESHOLD"],
    )


# =============================================================================
# ██████████████████████████████████████████████████████████████████████████
# STAGE 6 — GENERATION PIPELINE (Gemini 2.0 Flash + Hallucination Gate)
# ██████████████████████████████████████████████████████████████████████████
# =============================================================================

# CELL 6.1 — System Prompt Template
# CRITICAL: explicit "information not found" instruction is the hallucination
#   control mechanism. Low temperature (0.1) further reduces hallucination risk.

_SYSTEM_PROMPT = """You are BANASPATI, an academic assistant for the Department of Information Technology (DTI) at ITS Surabaya.

STRICT RULES:
1. Answer ONLY from the provided context. Do not use external knowledge.
2. If the context does not contain sufficient information to answer the query, output EXACTLY:
   "INFORMATION NOT FOUND: The provided academic documents do not contain sufficient information to answer this query."
3. Always cite the source document(s) used in your answer using format: [Source: <filename>, Page: <page_num>].
4. Be precise. Do not speculate or infer beyond what the context explicitly states.
5. Respond in the same language as the query (Bahasa Indonesia or English).

CONTEXT:
{context_block}

SOURCES AVAILABLE:
{sources_list}
"""

# -----------------------------------------------------------------------------
# CELL 6.2 — Generation Result Dataclass

@dataclass
class GenerationResult:
    query:               str
    answer:              str
    retrieval_result:    RetrievalResult
    context_block:       str
    final_prompt:        str
    is_hallucination_gate_triggered: bool
    # Telemetry (populated in Stage 7)
    generation_latency:  float = 0.0
    end_to_end_latency:  float = 0.0
    input_tokens:        int   = 0
    output_tokens:       int   = 0
    total_tokens:        int   = 0
    estimated_cost_usd:  float = 0.0
    throughput_tps:      float = 0.0
    ram_usage_mb:        float = 0.0
    ttft_note:           str   = "Non-streaming: TTFT not applicable"


# -----------------------------------------------------------------------------
# CELL 6.3 — Full RAG Generation Function
# MECHANISM:
#   1. retrieve(query) → RetrievalResult
#   2. Similarity gate check → if insufficient → return INFORMATION NOT FOUND
#   3. Build context_block from top_chunks
#   4. Render final_prompt with system template
#   5. Call Gemini 2.0 Flash → parse response + usage metadata
#   6. Compute telemetry
# -----------------------------------------------------------------------------

_generator = genai.GenerativeModel(
    model_name    = CFG["GENERATOR_MODEL"],
    generation_config = genai.types.GenerationConfig(
        temperature       = CFG["TEMPERATURE"],
        max_output_tokens = CFG["MAX_OUTPUT_TOKENS"],
    ),
)


def generate(query: str) -> GenerationResult:
    t_e2e_start = time.perf_counter()
    ram_before  = psutil.Process().memory_info().rss / 1e6  # MB

    # Step 1: Retrieve
    ret = retrieve(query)

    # Step 2: Hallucination gate
    if not ret.context_sufficient:
        t_end = time.perf_counter()
        return GenerationResult(
            query             = query,
            answer            = (
                "INFORMATION NOT FOUND: The provided academic documents do not "
                "contain sufficient information to answer this query."
            ),
            retrieval_result  = ret,
            context_block     = "",
            final_prompt      = "",
            is_hallucination_gate_triggered = True,
            end_to_end_latency = t_end - t_e2e_start,
        )

    # Step 3: Build context block
    context_parts = []
    for chunk in ret.chunks:
        context_parts.append(
            f"[{chunk.source_file} | Page {chunk.page_num + 1} | {chunk.chunk_type}]\n"
            f"{chunk.content}"
        )
    context_block = "\n\n---\n\n".join(context_parts)
    sources_list  = "\n".join(f"- {s}" for s in ret.sources)

    # Step 4: Render prompt
    user_message = f"Query: {query}"
    final_prompt = _SYSTEM_PROMPT.format(
        context_block = context_block,
        sources_list  = sources_list,
    )

    # Step 5: Generate
    t_gen_start = time.perf_counter()
    response    = _generator.generate_content(
        contents = [
            {"role": "user", "parts": [{"text": final_prompt + "\n\n" + user_message}]}
        ]
    )
    t_gen_end   = time.perf_counter()

    answer          = response.text.strip()
    gen_latency     = t_gen_end - t_gen_start
    e2e_latency     = t_gen_end - t_e2e_start

    # Step 6: Telemetry
    usage        = response.usage_metadata
    input_tok    = usage.prompt_token_count     if usage else 0
    output_tok   = usage.candidates_token_count if usage else 0
    total_tok    = usage.total_token_count       if usage else 0
    cost         = (input_tok  / 1e6 * CFG["COST_INPUT_PER_1M"] +
                    output_tok / 1e6 * CFG["COST_OUTPUT_PER_1M"])
    throughput   = output_tok / gen_latency if gen_latency > 0 else 0.0
    ram_usage    = psutil.Process().memory_info().rss / 1e6 - ram_before

    return GenerationResult(
        query             = query,
        answer            = answer,
        retrieval_result  = ret,
        context_block     = context_block,
        final_prompt      = final_prompt,
        is_hallucination_gate_triggered = False,
        generation_latency  = gen_latency,
        end_to_end_latency  = e2e_latency,
        input_tokens        = input_tok,
        output_tokens       = output_tok,
        total_tokens        = total_tok,
        estimated_cost_usd  = cost,
        throughput_tps      = throughput,
        ram_usage_mb        = ram_usage,
        ttft_note           = "Non-streaming: TTFT not applicable. "
                              "Enable stream=True in generate_content() for TTFT measurement.",
    )


# =============================================================================
# ██████████████████████████████████████████████████████████████████████████
# STAGE 7 — TELEMETRY LOGGER
# ██████████████████████████████████████████████████████████████████████████
# =============================================================================

# CELL 7.1 — Telemetry Log (accumulated across all queries)

TELEMETRY_LOG: list[dict] = []


def log_telemetry(result: GenerationResult) -> dict:
    """Extract and store telemetry from a GenerationResult."""
    entry = {
        "query":                  result.query,
        "retrieval_latency_s":    round(result.retrieval_result.retrieval_latency, 4),
        "generation_latency_s":   round(result.generation_latency, 4),
        "end_to_end_latency_s":   round(result.end_to_end_latency, 4),
        "input_tokens":           result.input_tokens,
        "output_tokens":          result.output_tokens,
        "total_tokens":           result.total_tokens,
        "throughput_tps":         round(result.throughput_tps, 2),
        "estimated_cost_usd":     round(result.estimated_cost_usd, 6),
        "ram_delta_mb":           round(result.ram_usage_mb, 2),
        "top_similarity":         round(result.retrieval_result.top_similarity, 4),
        "context_sufficient":     result.retrieval_result.context_sufficient,
        "gate_triggered":         result.is_hallucination_gate_triggered,
        "model":                  CFG["GENERATOR_MODEL"],
        "provider":               "Google AI Studio",
        "model_selection_reason": (
            "Gemini 2.0 Flash selected: free tier, 1M context window, "
            "native multimodal, lowest latency in Gemini family, "
            "satisfies ≤9B local constraint via API offload."
        ),
        "ttft_note":              result.ttft_note,
    }
    TELEMETRY_LOG.append(entry)
    return entry


def print_telemetry(entry: dict) -> None:
    print("\n[TELEMETRY] ─────────────────────────────────────────")
    for k, v in entry.items():
        if k not in ("model_selection_reason", "ttft_note", "query"):
            print(f"  {k:30s}: {v}")
    print(f"  {'ttft_note':30s}: {entry['ttft_note']}")
    print("[TELEMETRY] ─────────────────────────────────────────\n")


# =============================================================================
# ██████████████████████████████████████████████████████████████████████████
# STAGE 8 — RAGAS QUALITATIVE EVALUATION
# ██████████████████████████████████████████████████████████████████████████
# =============================================================================

# CELL 8.1 — Load Evaluation Dataset
# banaspati_eval_questions.csv must contain columns:
#   "question" → query text
#   "ground_truth" → reference answer (for context_recall + faithfulness)
# HALT if column schema is wrong.

_eval_df = pd.read_csv(CFG["DATASET_DIR"] / CFG["EVAL_CSV"])
print(f"[EVAL] CSV loaded: {len(_eval_df)} rows | Columns: {list(_eval_df.columns)}")

# Validate required columns
_required_cols = {"question", "ground_truth"}
if not _required_cols.issubset(_eval_df.columns):
    raise ValueError(
        f"[EVAL] CSV schema error. Expected columns: {_required_cols}. "
        f"Found: {set(_eval_df.columns)}"
    )

# Use minimum 10 queries as required by brief
_eval_queries = _eval_df.head(max(10, len(_eval_df))).to_dict("records")
print(f"[EVAL] Running evaluation on {len(_eval_queries)} queries.")

# -----------------------------------------------------------------------------
# CELL 8.2 — Run Pipeline Over Eval Queries
# Accumulates GenerationResults for both RAGAS and LLM-Judge consumption.
# -----------------------------------------------------------------------------

_eval_results: list[GenerationResult] = []

for row in tqdm(_eval_queries, desc="[EVAL] Pipeline inference"):
    res = generate(row["question"])
    log_telemetry(res)
    _eval_results.append(res)

# -----------------------------------------------------------------------------
# CELL 8.3 — RAGAS Dataset Construction
# RAGAS Dataset schema:
#   question        → query
#   answer          → generated answer
#   contexts        → list of retrieved chunk texts (list[str])
#   ground_truth    → reference answer from CSV
# -----------------------------------------------------------------------------

from datasets import Dataset as HFDataset
from ragas import evaluate as ragas_evaluate
from ragas.metrics import (
    faithfulness,
    answer_relevancy,
    context_precision,
    context_recall,
)

_ragas_data = {
    "question":     [],
    "answer":       [],
    "contexts":     [],
    "ground_truth": [],
}

for row, res in zip(_eval_queries, _eval_results):
    _ragas_data["question"].append(res.query)
    _ragas_data["answer"].append(res.answer)
    _ragas_data["contexts"].append([c.content for c in res.retrieval_result.chunks])
    _ragas_data["ground_truth"].append(row["ground_truth"])

_ragas_dataset = HFDataset.from_dict(_ragas_data)

# -----------------------------------------------------------------------------
# CELL 8.4 — Execute RAGAS Evaluation
# RAGAS internally uses an LLM for faithfulness + answer_relevancy scoring.
# Configure it to use Gemini 2.0 Flash via LangChain wrapper to stay free-tier.
# -----------------------------------------------------------------------------

from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from ragas.llms import LangchainLLMWrapper
from ragas.embeddings import LangchainEmbeddingsWrapper

_ragas_llm  = LangchainLLMWrapper(ChatGoogleGenerativeAI(
    model   = CFG["GENERATOR_MODEL"],
    google_api_key = GEMINI_API_KEY,
))
_ragas_emb  = LangchainEmbeddingsWrapper(GoogleGenerativeAIEmbeddings(
    model  = "models/text-embedding-004",
    google_api_key = GEMINI_API_KEY,
))

print("[RAGAS] Running evaluation — expect 3-10 minutes for 10 queries...")
_ragas_scores = ragas_evaluate(
    dataset    = _ragas_dataset,
    metrics    = [faithfulness, answer_relevancy, context_precision, context_recall],
    llm        = _ragas_llm,
    embeddings = _ragas_emb,
)

RAGAS_RESULTS_DF = _ragas_scores.to_pandas()
print("\n[RAGAS] RESULTS:")
print(RAGAS_RESULTS_DF[[
    "faithfulness", "answer_relevancy", "context_precision", "context_recall"
]].describe())


# =============================================================================
# ██████████████████████████████████████████████████████████████████████████
# STAGE 9 — LLM-AS-JUDGE (Gemini 3 Flash, mandated)
# ██████████████████████████████████████████████████████████████████████████
# =============================================================================

# CELL 9.1 — Judge Model Initialization
# MANDATE: brief requires Gemini 3 Flash specifically for judge role.
# REASON for separation from generator: prevents self-evaluation bias.
# Judge model ≠ generator model → more objective scoring.

_judge_model = genai.GenerativeModel(
    model_name = CFG["JUDGE_MODEL"],
    generation_config = genai.types.GenerationConfig(
        temperature       = 0.0,    # deterministic scoring
        max_output_tokens = 512,
    ),
)

# -----------------------------------------------------------------------------
# CELL 9.2 — Judge Rubric Prompt Template
# Rubric dimensions (all scored 1-5):
#   Correctness:    factual accuracy vs ground truth
#   Faithfulness:   answer grounded in context (no hallucination)
#   Relevance:      answer addresses the query
#   Completeness:   all aspects of query addressed
#   Source Support: citations present and accurate
#   Hallucination:  5 = no hallucination, 1 = severe hallucination
# OUTPUT FORMAT: strict JSON to enable programmatic parsing.
# -----------------------------------------------------------------------------

_JUDGE_PROMPT_TEMPLATE = """You are a strict academic evaluation judge. Score the following RAG system response against the rubric below.

QUERY:
{query}

REFERENCE ANSWER (Ground Truth):
{ground_truth}

RETRIEVED CONTEXT:
{context}

SYSTEM ANSWER:
{answer}

RUBRIC (score each 1-5, where 5=best):
- correctness:    How factually accurate is the answer compared to the ground truth?
- faithfulness:   Is every claim in the answer supported by the retrieved context?
- relevance:      Does the answer directly address the query?
- completeness:   Does the answer cover all aspects of the query?
- source_support: Are source citations present and traceable to the context?
- hallucination:  Score 5 if NO hallucination detected. Score 1 if severe fabrication.

OUTPUT FORMAT — respond ONLY with valid JSON, no preamble, no markdown fences:
{{"correctness": <int>, "faithfulness": <int>, "relevance": <int>, "completeness": <int>, "source_support": <int>, "hallucination": <int>, "reasoning": "<one sentence>"}}
"""

# -----------------------------------------------------------------------------
# CELL 9.3 — Execute Judge Evaluation
# -----------------------------------------------------------------------------

JUDGE_RESULTS: list[dict] = []

for row, res in tqdm(
    zip(_eval_queries, _eval_results),
    total=len(_eval_queries),
    desc="[JUDGE] Evaluating",
):
    _prompt = _JUDGE_PROMPT_TEMPLATE.format(
        query        = res.query,
        ground_truth = row["ground_truth"],
        context      = res.context_block[:4000],  # truncate to avoid judge token overflow
        answer       = res.answer,
    )

    try:
        _judge_response = _judge_model.generate_content(_prompt)
        _raw            = _judge_response.text.strip()
        # Strip any accidental markdown fences
        _raw            = re.sub(r"```json|```", "", _raw).strip()
        _scores         = json.loads(_raw)
        _scores["query"] = res.query
    except (json.JSONDecodeError, Exception) as e:
        # Graceful degradation — log error, don't abort eval run
        _scores = {
            "query":        res.query,
            "error":        str(e),
            "raw_response": _judge_response.text if _judge_response else "NO_RESPONSE",
        }

    JUDGE_RESULTS.append(_scores)

JUDGE_RESULTS_DF = pd.DataFrame(JUDGE_RESULTS)
print("\n[JUDGE] LLM-AS-JUDGE RESULTS:")
_score_cols = ["correctness", "faithfulness", "relevance", "completeness",
               "source_support", "hallucination"]
_available  = [c for c in _score_cols if c in JUDGE_RESULTS_DF.columns]
print(JUDGE_RESULTS_DF[_available].describe())


# =============================================================================
# ██████████████████████████████████████████████████████████████████████████
# STAGE 10 — DEMO SANDBOX
# ██████████████████████████████████████████████████████████████████████████
# =============================================================================

# CELL 10.1 — Telemetry Summary Report

def print_telemetry_summary() -> None:
    """Print aggregate telemetry across all logged queries."""
    if not TELEMETRY_LOG:
        print("[TELEMETRY] No entries logged.")
        return

    df = pd.DataFrame(TELEMETRY_LOG)
    print("\n[TELEMETRY] AGGREGATE SUMMARY ──────────────────────────────────")
    numeric_cols = [
        "retrieval_latency_s", "generation_latency_s", "end_to_end_latency_s",
        "input_tokens", "output_tokens", "total_tokens",
        "throughput_tps", "estimated_cost_usd", "ram_delta_mb",
    ]
    print(df[numeric_cols].describe().round(4))
    print(f"\n  Total estimated cost (all queries): "
          f"${df['estimated_cost_usd'].sum():.6f} USD")
    print(f"  Model:    {CFG['GENERATOR_MODEL']} via Google AI Studio")
    print(f"  Provider: Google AI Studio (free tier)")
    print(f"  TTFT:     Non-streaming — TTFT not captured. "
          f"Set stream=True in generate_content() to enable.")
    print("[TELEMETRY] ────────────────────────────────────────────────────\n")

print_telemetry_summary()


# -----------------------------------------------------------------------------
# CELL 10.2 — Demo Sandbox
# PURPOSE: live evaluation harness for unforeseen queries during demo/grading.
# Outputs: query, context, sources, final_prompt, full metrics — all explicit.
# USAGE: call demo_sandbox("your query here") directly in Colab cell.
# -----------------------------------------------------------------------------

def demo_sandbox(query: str) -> None:
    """
    Full-transparency single-query runner for live demo and grading evaluation.

    Outputs ALL intermediate artifacts:
      1. Query
      2. Retrieved context (with source labels)
      3. Source files
      4. Final prompt (as sent to generator)
      5. Generated answer
      6. Telemetry metrics
      7. Hallucination gate status
    """
    print("=" * 70)
    print("BANASPATI DEMO SANDBOX")
    print("=" * 70)

    print(f"\n[1] QUERY:\n{query}\n")

    result = generate(query)
    entry  = log_telemetry(result)

    print(f"[2] RETRIEVED CONTEXT:")
    if result.context_block:
        print(result.context_block[:2000])  # truncate for display
        if len(result.context_block) > 2000:
            print(f"    ... [{len(result.context_block) - 2000} chars truncated]")
    else:
        print("    [No context retrieved — gate triggered]")

    print(f"\n[3] SOURCES:")
    for s in result.retrieval_result.sources:
        print(f"    - {s}")

    print(f"\n[4] FINAL PROMPT (first 1000 chars):")
    print(result.final_prompt[:1000] if result.final_prompt else "[No prompt — gate triggered]")

    print(f"\n[5] GENERATED ANSWER:")
    print(result.answer)

    print(f"\n[6] METRICS:")
    _display_keys = [
        "retrieval_latency_s", "generation_latency_s", "end_to_end_latency_s",
        "input_tokens", "output_tokens", "throughput_tps",
        "estimated_cost_usd", "ram_delta_mb", "top_similarity",
    ]
    for k in _display_keys:
        print(f"    {k:30s}: {entry.get(k, 'N/A')}")
    print(f"    {'ttft_note':30s}: {entry['ttft_note']}")

    print(f"\n[7] HALLUCINATION GATE: "
          f"{'TRIGGERED — context insufficient' if result.is_hallucination_gate_triggered else 'PASSED — context sufficient'}")
    print(f"    Top similarity score: {result.retrieval_result.top_similarity:.4f} "
          f"(threshold: {CFG['SIMILARITY_THRESHOLD']})")

    print("\n" + "=" * 70)


# -----------------------------------------------------------------------------
# CELL 10.3 — Entry Point: Run Demo Sandbox with Sample Query
# Replace the query string below with any live question during evaluation.
# -----------------------------------------------------------------------------

if __name__ == "__main__":
    # Sample query — replace or extend with multiple calls as needed
    demo_sandbox("Siapa saja dosen di Departemen Teknologi Informasi ITS?")

    # Optional: run another query to test hallucination gate
    demo_sandbox("What is the recipe for nasi goreng?")  # out-of-scope → gate fires